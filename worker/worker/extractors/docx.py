from __future__ import annotations
from io import BytesIO

from docx import Document

from . import NormalizedDoc, PageContent


def extract(source_id: str, raw: bytes) -> NormalizedDoc:
    doc = Document(BytesIO(raw))
    section_title = None
    buf: list[str] = []
    tables: list[str] = []

    for p in doc.paragraphs:
        text = p.text or ""
        if not text.strip():
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if style.startswith("heading"):
            section_title = text.strip()
        buf.append(text)

    for t in doc.tables:
        rows = []
        for row in t.rows:
            rows.append("\t".join(c.text.strip() for c in row.cells))
        tables.append("\n".join(rows))

    return NormalizedDoc(
        source_id=source_id,
        pages=[
            PageContent(
                page_num=1,
                text="\n".join(buf),
                tables=tables,
                section_title=section_title,
            )
        ],
    )
